# import frappe
# import os
# from pdfminer.high_level import extract_text
# from frappe.model.document import Document
# import pytesseract
# from pdf2image import convert_from_path
# import google.generativeai as genai
# import json

# class PDFUpload(Document):
#     pass

# @frappe.whitelist()
# def process_pdfs(docname):
#     """Enqueues the PDF processing job."""
#     frappe.enqueue("resume.resume.doctype.pdf_upload.pdf_upload.process_pdfs_background", queue="long", docname=docname)
#     frappe.msgprint("PDF processing has been queued in the background.")

# @frappe.whitelist()
# def list_available_gemini_models():
#     """List available Gemini models for debugging."""
#     try:
#         api_key = frappe.conf.get("gemini_api_key")
#         if not api_key:
#             return {"error": "Gemini API key missing in site_config.json"}
        
#         genai.configure(api_key=api_key)
#         models = genai.list_models()
#         available_models = []
#         model_details = []
#         for model in models:
#             if 'generateContent' in model.supported_generation_methods:
#                 model_name = model.name
#                 # Extract just the model identifier (remove 'models/' prefix if present)
#                 model_id = model_name.replace('models/', '')
#                 available_models.append(model_id)
#                 model_details.append({
#                     "full_name": model_name,
#                     "model_id": model_id,
#                     "display_name": getattr(model, 'display_name', 'N/A'),
#                     "supported_methods": list(model.supported_generation_methods)
#                 })
        
#         # Log to frappe.log for easy viewing
#         frappe.logger().info(f"Available Gemini models: {available_models}")
#         for detail in model_details:
#             frappe.logger().info(f"Model: {detail['model_id']} (Full: {detail['full_name']})")
        
#         return {
#             "available_models": available_models,
#             "model_details": model_details
#         }
#     except Exception as e:
#         error_msg = str(e)
#         frappe.logger().error(f"Error listing Gemini models: {error_msg}")
#         return {"error": error_msg}

# def parse_with_gemini(text, job_title=None, job_description=None):
#     """Send resume text to Gemini model and return structured JSON."""
#     try:
#         api_key = frappe.conf.get("gemini_api_key")
#         if not api_key:
#             raise Exception("Gemini API key missing in site_config.json")

#         genai.configure(api_key=api_key)

#         # Load prompt template
#         prompt_path = frappe.get_app_path("resume", "resume", "doctype", "pdf_upload", "resume_prompt.txt")
#         with open(prompt_path, "r") as f:
#             prompt_template = f.read()

#         prompt = prompt_template.replace("{{RESUME_TEXT}}", text)
        
#         if job_title:
#             prompt = prompt.replace("{{JOB_TITLE}}", job_title)
#         else:
#             prompt = prompt.replace("{{JOB_TITLE}}", "N/A")
            
#         if job_description:
#             prompt = prompt.replace("{{JOB_DESCRIPTION}}", job_description)
#         else:
#             prompt = prompt.replace("{{JOB_DESCRIPTION}}", "N/A")

#         # Use a hardcoded list of prioritized models to avoid latency of listing models
#         model_names = [
#             "gemini-1.5-flash",
#             "gemini-1.5-pro",
#             "gemini-pro",
#         ]
        
#         last_error = None
#         skipped_models = []
        
#         # Helper function to try a list of models
#         def try_models(models_to_try):
#             nonlocal last_error
#             for model_name in models_to_try:
#                 try:
#                     frappe.logger().info(f"Attempting to use Gemini model: {model_name}")
#                     model = genai.GenerativeModel(model_name)
#                     response = model.generate_content(prompt)
#                     cleaned = response.text.strip()
#                     # Remove markdown code blocks if present
#                     if cleaned.startswith("```json"):
#                         cleaned = cleaned[7:]
#                     if cleaned.endswith("```"):
#                         cleaned = cleaned[:-3]
#                     cleaned = cleaned.strip()
                    
#                     applicant_data = json.loads(cleaned)
#                     frappe.logger().info(f"Successfully used Gemini model: {model_name}")
#                     return applicant_data
#                 except Exception as e:
#                     error_str = str(e)
#                     # Skip models that require Computer Use tool or other special tools
#                     if "Computer Use tool" in error_str or "requires the use of" in error_str:
#                         skipped_models.append(model_name)
#                         frappe.logger().warning(f"Skipping model {model_name} (requires special tools: {error_str[:100]})")
#                         continue
#                     last_error = e
#                     frappe.logger().warning(f"Failed to use model {model_name}: {error_str[:200]}")
#                     continue
#             return None

#         # 1. Try hardcoded models first
#         result = try_models(model_names)
#         if result:
#             return result
            
#         # 2. If hardcoded models failed, try listing available models dynamically
#         frappe.logger().warning("All hardcoded Gemini models failed. Attempting to list available models dynamically.")
#         try:
#             models = genai.list_models()
#             dynamic_models = []
#             excluded_keywords = ['computer', 'agent', 'function-calling']
            
#             for model in models:
#                 if 'generateContent' in model.supported_generation_methods:
#                     model_id = model.name.replace('models/', '')
#                     model_lower = model_id.lower()
#                     if not any(keyword in model_lower for keyword in excluded_keywords):
#                         dynamic_models.append(model_id)
            
#             if dynamic_models:
#                 frappe.logger().info(f"Found dynamic models: {dynamic_models}")
#                 # Filter out ones we already tried to avoid duplicate work/logs
#                 new_models = [m for m in dynamic_models if m not in model_names]
#                 if new_models:
#                     result = try_models(new_models)
#                     if result:
#                         return result
#         except Exception as e:
#             frappe.logger().error(f"Failed to list dynamic models: {e}")
        
#         if skipped_models:
#             frappe.logger().info(f"Skipped {len(skipped_models)} models that require special tools: {skipped_models}")
        
#         # If all models failed, raise the last error
#         if last_error:
#             raise Exception(f"All Gemini models failed. Last error: {last_error}")

#     except Exception as e:
#         frappe.logger().error(f"Gemini parsing failed: {e}")
#         raise

# def process_pdfs_background(docname):
#     """Reads uploaded PDF files, extracts text, and creates Job Applicant records."""
#     try:
#         doc = frappe.get_doc("PDF Upload", docname)

#         if not doc.get("files"):
#             frappe.logger().warning(f"PDF Upload {docname}: No files found.")
#             return

#         for file_entry in doc.get("files"):
#             file_url = file_entry.get("file_upload")

#             if not file_url:
#                 frappe.logger().error(f"PDF Upload {docname}: Missing file URL.")
#                 continue
            
#             # Determine if the file is public or private
#             if file_url.startswith("/private/files/"):
#                 file_path = frappe.get_site_path(file_url.lstrip("/"))
#             elif file_url.startswith("/files/"):
#                 file_path = frappe.get_site_path("public" + file_url)
#             else:
#                 frappe.logger().error(f"PDF Upload {docname}: Invalid file path {file_url}")
#                 continue

#             # Validate file existence
#             if not os.path.exists(file_path):
#                 frappe.logger().error(f"PDF Upload {docname}: File not found on server {file_path}")
#                 continue

#             try:
#                 # Extract text using pdfminer
#                 text = extract_text(file_path)
                
#                 # If text extraction failed, try using OCR
#                 if not text or not text.strip():
#                     text = extract_text_with_ocr(file_path)
#             except Exception as e:
#                 frappe.logger().error(f"PDF Upload {docname}: Error extracting text from {file_url}: {e}")
#                 continue

#             # Parse extracted text
#             try:
#                 job_description = None
#                 if doc.job_title:
#                     job_description = frappe.db.get_value("Job Opening", doc.job_title, "description")
                
#                 applicant_data = parse_with_gemini(text, doc.job_title, job_description)
#             except Exception as e:
#                  frappe.logger().error(f"PDF Upload {docname}: Error parsing text from {file_url}: {e}")
#                  continue
            
#             if not applicant_data.get("applicant_name") or not applicant_data.get("email_id"):
#                 frappe.logger().warning(f"PDF Upload {docname}: Missing required fields (name/email) in {file_url}")
#                 continue

#             # Check if applicant already exists using email
#             if frappe.db.exists("Job Applicant", {"email_id": applicant_data["email_id"]}):
#                 frappe.logger().info(f"PDF Upload {docname}: Applicant with email {applicant_data['email_id']} already exists.")
#                 continue

#             # Create new Job Applicant record
#             try:
#                 # Format notes with concise scoring details (max 140 chars)
#                 fit_level = applicant_data.get('fit_level', 'N/A')
#                 score = applicant_data.get('score', 'N/A')
#                 rating = applicant_data.get('applicant_rating', 'N/A')
#                 applicant = frappe.get_doc({
#                     "doctype": "Job Applicant",
#                     "applicant_name": applicant_data.get("applicant_name"),
#                     "email_id": applicant_data.get("email_id"),
#                     "phone_number": applicant_data.get("phone_number"),
#                     "resume_attachment": file_url,
#                     "status": "Open",
#                     "job_title": doc.job_title,
#                     "designation": doc.designation,
#                     "applicant_rating": applicant_data.get("applicant_rating"),
#                     "score": score,
#                     "fit_level": fit_level,
#                     "justification_by_ai": applicant_data.get('justification_by_ai', 'N/A')
#                 })
#                 applicant.insert(ignore_permissions=True)
#                 frappe.logger().info(f"PDF Upload {docname}: Created Job Applicant {applicant.name}")
#             except Exception as e:
#                 frappe.logger().error(f"PDF Upload {docname}: Failed to create applicant for {file_url}: {e}")

#     except Exception as e:
#         frappe.logger().error(f"PDF Upload {docname}: Critical error in background job: {e}")

# def extract_text_with_ocr(file_path):
#     """Extract text from scanned PDFs using Tesseract OCR."""
#     try:
#         images = convert_from_path(file_path)
#         extracted_text = ""
#         for image in images:
#             extracted_text += pytesseract.image_to_string(image)
#         return extracted_text.strip()
#     except Exception as e:
#         raise ValueError(f"Error extracting text with OCR: {e}")








# 2

# import frappe
# import os
# import json
# import mimetypes
# from pdfminer.high_level import extract_text
# from frappe.model.document import Document
# import pytesseract
# from pdf2image import convert_from_path
# from PIL import Image
# import docx  # For .docx files
# import google.generativeai as genai

# class PDFUpload(Document):
#     pass

# @frappe.whitelist()
# def process_pdfs(docname):
#     """Enqueues the processing job for all types of files."""
#     frappe.enqueue("resume.resume.doctype.pdf_upload.pdf_upload.process_files_background", queue="long", docname=docname)
#     frappe.msgprint("File processing has been queued in the background.")

# # --- AI Logic (UNCHANGED as per request) ---
# @frappe.whitelist()
# def list_available_gemini_models():
#     """List available Gemini models for debugging."""
#     try:
#         api_key = frappe.conf.get("gemini_api_key")
#         if not api_key:
#             return {"error": "Gemini API key missing in site_config.json"}
        
#         genai.configure(api_key=api_key)
#         models = genai.list_models()
#         available_models = []
#         model_details = []
#         for model in models:
#             if 'generateContent' in model.supported_generation_methods:
#                 model_name = model.name
#                 model_id = model_name.replace('models/', '')
#                 available_models.append(model_id)
#                 model_details.append({
#                     "full_name": model_name,
#                     "model_id": model_id,
#                     "display_name": getattr(model, 'display_name', 'N/A'),
#                     "supported_methods": list(model.supported_generation_methods)
#                 })
        
#         frappe.logger().info(f"Available Gemini models: {available_models}")
#         return {"available_models": available_models, "model_details": model_details}
#     except Exception as e:
#         frappe.logger().error(f"Error listing Gemini models: {str(e)}")
#         return {"error": str(e)}

# def parse_with_gemini(text, job_title=None, job_description=None):
#     """AI Logic: Send resume text to Gemini model and return structured JSON."""
#     # (No changes made to this function as requested)
#     try:
#         api_key = frappe.conf.get("gemini_api_key")
#         if not api_key:
#             raise Exception("Gemini API key missing in site_config.json")

#         genai.configure(api_key=api_key)
#         prompt_path = frappe.get_app_path("resume", "resume", "doctype", "pdf_upload", "resume_prompt.txt")
#         with open(prompt_path, "r") as f:
#             prompt_template = f.read()

#         prompt = prompt_template.replace("{{RESUME_TEXT}}", text)
#         prompt = prompt.replace("{{JOB_TITLE}}", job_title or "N/A")
#         prompt = prompt.replace("{{JOB_DESCRIPTION}}", job_description or "N/A")

#         model_names = ["gemini-1.5-flash", "gemini-1.5-pro", "gemini-pro"]
#         last_error = None
#         skipped_models = []
        
#         def try_models(models_to_try):
#             nonlocal last_error
#             for model_name in models_to_try:
#                 try:
#                     frappe.logger().info(f"Attempting Gemini model: {model_name}")
#                     model = genai.GenerativeModel(model_name)
#                     response = model.generate_content(prompt)
#                     cleaned = response.text.strip()
#                     if cleaned.startswith("```json"): cleaned = cleaned[7:]
#                     if cleaned.endswith("```"): cleaned = cleaned[:-3]
#                     return json.loads(cleaned.strip())
#                 except Exception as e:
#                     if "Computer Use tool" in str(e):
#                         skipped_models.append(model_name)
#                         continue
#                     last_error = e
#                     continue
#             return None

#         result = try_models(model_names)
#         if result: return result
        
#         # Dynamic fallback
#         models = genai.list_models()
#         dynamic_models = [m.name.replace('models/', '') for m in models 
#                           if 'generateContent' in m.supported_generation_methods 
#                           and m.name.replace('models/', '') not in model_names]
#         return try_models(dynamic_models) or (lambda: exec('raise last_error'))()
#     except Exception as e:
#         frappe.logger().error(f"Gemini parsing failed: {e}")
#         raise

# # --- Text Extraction Logic (Enhanced for all formats) ---

# def extract_text_from_any_file(file_path):
#     """Detects file type and extracts text accordingly."""
#     ext = os.path.splitext(file_path)[1].lower()
#     text = ""

#     frappe.logger().info(f"Starting extraction for file: {file_path} with extension: {ext}")

#     if ext == ".pdf":
#         text = extract_text(file_path)
#         if not text.strip():
#             # frappe.logger().info(f"PDF looks like an image, starting OCR: {file_path}")
#             frappe.log_error(
#                 title="trigger_ai_call file_path",
#                 message=f"file_path: {file_path}"
#             )
#             text = extract_text_with_ocr(file_path)
    
#     elif ext == ".docx":
#         doc = docx.Document(file_path)
#         text = "\n".join([para.text for para in doc.paragraphs])
    
#     elif ext in [".jpg", ".jpeg", ".png"]:
#         text = pytesseract.image_to_string(Image.open(file_path))
    
#     elif ext == ".txt":
#         with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
#             text = f.read()
    
#     else:
#         frappe.logger().warning(f"Unsupported file format: {ext}")
#         return None

#     return text.strip()

# def extract_text_with_ocr(file_path):
#     """OCR for Scanned PDFs."""
#     try:
#         images = convert_from_path(file_path)
#         return "\n".join([pytesseract.image_to_string(img) for img in images])
#     except Exception as e:
#         frappe.logger().error(f"OCR Error: {str(e)}")
#         return ""

# import base64
# def parse_with_gemini_file(file_path, job_title=None, job_description=None):
#     api_key = frappe.conf.get("gemini_api_key")
#     if not api_key:
#         raise Exception("Gemini API key missing in site_config.json")

#     genai.configure(api_key=api_key)
#     # model = get_gemini()
#     model = genai.GenerativeModel("gemini-2.5-pro")

#     with open(file_path, "rb") as f:
#         pdf_bytes = f.read()
        
#     prompt_path = frappe.get_app_path("resume", "resume", "doctype", "pdf_upload", "resume_prompt.txt")
#     with open(prompt_path, "r") as f:
#         prompt_template = f.read()

#     prompt = prompt_template.replace("{{RESUME_TEXT}}", "Resume attached as PDF.")
#     prompt = prompt.replace("{{JOB_TITLE}}", job_title or "N/A")
#     prompt = prompt.replace("{{JOB_DESCRIPTION}}", job_description or "N/A")

#     # prompt = PROMPT

#     response = model.generate_content(
#         [
#             {"mime_type": "application/pdf", "data": pdf_bytes},
#             prompt
#         ]
#     )

#     text = response.text.strip()

#     if text.startswith("```"):
#         text = text.replace("```json", "").replace("```", "").strip()

#     return json.loads(text)

# # --- Background Job (Enhanced Logging) ---

# def process_files_background(docname):
#     """Processes any uploaded file, extracts text, and creates Applicants."""
#     try:
#         doc = frappe.get_doc("PDF Upload", docname)
#         if not doc.get("files"):
#             frappe.logger().warning(f"PDF Upload {docname}: No files found.")
#             return

#         for file_entry in doc.get("files"):
#             file_url = file_entry.get("file_upload")
#             if not file_url: continue
            
#             # File Path logic
#             if file_url.startswith("/private/files/"):
#                 file_path = frappe.get_site_path(file_url.lstrip("/"))
#             else:
#                 file_path = frappe.get_site_path("public" + file_url)

#             if not os.path.exists(file_path):
#                 frappe.logger().error(f"File not found on server: {file_path}")
#                 continue

#             # 1. Extract Text
#             try:
#                 text = extract_text_from_any_file(file_path)
#                 if not text:
#                     frappe.logger().error(f"Could not extract any text from {file_url}")
#                     continue
#             except Exception as e:
#                 frappe.logger().error(f"Extraction error for {file_url}: {str(e)}")
#                 continue

#             # 2. Parse with AI
#             try:
#                 job_desc = frappe.db.get_value("Job Opening", doc.job_title, "description") if doc.job_title else None
#                 applicant_data = parse_with_gemini(text, doc.job_title, job_desc)
#             except Exception as e:
#                  frappe.logger().error(f"AI Parsing failed for {file_url}: {str(e)}")
#                  continue
            
#             # 3. Validation & Creation
#             email = applicant_data.get("email_id")
#             if not email or not applicant_data.get("applicant_name"):
#                 frappe.logger().warning(f"Missing Name/Email in AI response for {file_url}")
#                 continue

#             if frappe.db.exists("Job Applicant", {"email_id": email}):
#                 frappe.logger().info(f"Skipping: Applicant {email} already exists.")
#                 continue

#             try:
#                 new_applicant = frappe.get_doc({
#                     "doctype": "Job Applicant",
#                     "applicant_name": applicant_data.get("applicant_name"),
#                     "email_id": email,
#                     "phone_number": applicant_data.get("phone_number"),
#                     # "custom_phone_number_2": applicant_data.get("custom_phone_number_2", ""),  # ADD THIS
#                     "resume_attachment": file_url,
#                     "status": "Open",
#                     "job_title": doc.job_title,
#                     "designation": doc.designation,
#                     "applicant_rating": applicant_data.get("applicant_rating"),
#                     "score": applicant_data.get('score'),
#                     "fit_level": applicant_data.get('fit_level'),
#                     "justification_by_ai": applicant_data.get('justification_by_ai')
#                 })
#                 new_applicant.insert(ignore_permissions=True)
#                 frappe.logger().info(f"Successfully Created Applicant: {new_applicant.name} from {file_url}")
#             except Exception as e:
#                 frappe.logger().error(f"Database insertion failed for {email}: {str(e)}")

#     except Exception as e:
#         frappe.logger().error(f"Critical error in background job {docname}: {str(e)}")



# 3

import frappe
import os
import json
import mimetypes
import shutil
import subprocess
import time
from urllib.parse import urlparse
from concurrent.futures import ThreadPoolExecutor, as_completed
from pdfminer.high_level import extract_text
from frappe import _
from frappe.model.document import Document
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import docx  # For .docx files
import google.generativeai as genai

# Max concurrent Gemini API calls - keep low to avoid rate limits (429)
PARALLEL_WORKERS = 2
GEMINI_RETRY_ATTEMPTS = 4
GEMINI_RETRY_BASE_DELAY = 3  # seconds

class PDFUpload(Document):
    pass


def normalize_attach_file_url(file_url):
    """Attach fields may store a full URL; File lookup and disk paths need /files/... or /private/... paths."""
    if not file_url:
        return file_url
    fu = (file_url or "").strip()
    if fu.startswith(("http://", "https://")):
        fu = urlparse(fu).path or fu
    site = (frappe.utils.get_url() or "").rstrip("/")
    if site and fu.startswith(site + "/"):
        fu = fu[len(site) :]
    elif site and fu == site:
        fu = "/"
    if not fu.startswith("/"):
        fu = "/" + fu.lstrip("/")
    return fu


@frappe.whitelist()
def process_pdfs(docname):
    """Validates configuration and enqueues resume parsing (runs in a background worker)."""
    if not docname:
        frappe.throw(_("Document name is required."))
    if not frappe.db.exists("PDF Upload", docname):
        frappe.throw(_("PDF Upload {0} was not found.").format(frappe.bold(docname)))

    if not frappe.conf.get("gemini_api_key"):
        frappe.throw(
            _(
                "Set <code>gemini_api_key</code> in site_config.json. Without it, resumes cannot be parsed."
            ),
            title=_("Gemini API not configured"),
        )

    doc = frappe.get_doc("PDF Upload", docname)
    rows = doc.get("files") or []
    if not rows or not any(r.get("file_upload") for r in rows):
        frappe.throw(
            _("Add at least one file in the Files table, save, then run processing again."),
            title=_("No files to process"),
        )

    if doc.job_title:
        jo_status = frappe.db.get_value("Job Opening", doc.job_title, "status")
        if jo_status == "Closed":
            frappe.throw(
                _(
                    "Job Opening <b>{0}</b> is <b>Closed</b>. HR does not allow new Job Applicants "
                    "against closed openings. Reopen this Job Opening or select an active one."
                ).format(doc.job_title),
                title=_("Job Opening is closed"),
            )

    # default queue + long timeout: many servers only run `bench worker` (all queues) or default-only;
    # explicit timeout avoids RQ killing slow Gemini batches (default queue timeout is often 300s).
    frappe.enqueue(
        "resume.resume.doctype.pdf_upload.pdf_upload.process_files_background",
        queue="default",
        timeout=1800,
        docname=docname,
    )
    frappe.msgprint(
        _(
            "Resume processing has been queued. If nothing changes after a few minutes, confirm "
            "<code>bench worker</code> is running and check Error Log / logs/worker.error.log."
        ),
        title=_("Queued"),
        indicator="green",
    )

# --- AI Logic (UNCHANGED as per request) ---
@frappe.whitelist()
def list_available_gemini_models():
    """List available Gemini models for debugging."""
    try:
        api_key = frappe.conf.get("gemini_api_key")
        if not api_key:
            return {"error": "Gemini API key missing in site_config.json"}
        
        genai.configure(api_key=api_key)
        models = genai.list_models()
        available_models = []
        model_details = []
        for model in models:
            if 'generateContent' in model.supported_generation_methods:
                model_name = model.name
                model_id = model_name.replace('models/', '')
                available_models.append(model_id)
                model_details.append({
                    "full_name": model_name,
                    "model_id": model_id,
                    "display_name": getattr(model, 'display_name', 'N/A'),
                    "supported_methods": list(model.supported_generation_methods)
                })
        
        frappe.logger().info(f"Available Gemini models: {available_models}")
        return {"available_models": available_models, "model_details": model_details}
    except Exception as e:
        frappe.logger().error(f"Error listing Gemini models: {str(e)}")
        return {"error": str(e)}

def parse_with_gemini(text, job_title=None, job_description=None):
    """AI Logic: Send resume text to Gemini model and return structured JSON."""
    # (No changes made to this function as requested)
    try:
        api_key = frappe.conf.get("gemini_api_key")
        if not api_key:
            raise Exception("Gemini API key missing in site_config.json")

        genai.configure(api_key=api_key)
        prompt_path = frappe.get_app_path("resume", "resume", "doctype", "pdf_upload", "resume_prompt.txt")
        with open(prompt_path, "r") as f:
            prompt_template = f.read()

        prompt = prompt_template.replace("{{RESUME_TEXT}}", text)
        prompt = prompt.replace("{{JOB_TITLE}}", job_title or "N/A")
        prompt = prompt.replace("{{JOB_DESCRIPTION}}", job_description or "N/A")

        model_names = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-2.5-pro"]  # flash first for speed
        last_error = None
        skipped_models = []
        
        def try_models(models_to_try):
            nonlocal last_error
            for model_name in models_to_try:
                try:
                    frappe.logger().info(f"Gemini model: {model_name}")
                    model = genai.GenerativeModel(model_name)
                    response = model.generate_content(prompt)
                    cleaned = response.text.strip()
                    if cleaned.startswith("```json"): cleaned = cleaned[7:]
                    if cleaned.endswith("```"): cleaned = cleaned[:-3]
                    return json.loads(cleaned.strip())
                except Exception as e:
                    if "Computer Use tool" in str(e):
                        skipped_models.append(model_name)
                        continue
                    last_error = e
                    continue
            return None

        result = try_models(model_names)
        if result: return result
        
        # Dynamic fallback
        models = genai.list_models()
        dynamic_models = [m.name.replace('models/', '') for m in models 
                          if 'generateContent' in m.supported_generation_methods 
                          and m.name.replace('models/', '') not in model_names]
        return try_models(dynamic_models) or (lambda: exec('raise last_error'))()
    except Exception as e:
        frappe.logger().error(f"Gemini parsing failed: {e}")
        raise

# --- Text Extraction Logic (Enhanced for all formats) ---

def extract_text_from_legacy_doc(file_path):
    """Plain text from legacy Word .doc (OLE). python-docx does not read .doc.

    Uses doc2txt (bundled antiword on Linux amd64 / Windows / macOS ARM). Falls back to
    system ``antiword`` if the bundled binary is unavailable (e.g. some architectures).
    """
    try:
        from doc2txt import extract_text as doc_bin_extract

        return (doc_bin_extract(file_path) or "").strip()
    except Exception:
        pass
    antiword = shutil.which("antiword")
    if antiword:
        try:
            r = subprocess.run(
                [antiword, file_path],
                capture_output=True,
                text=True,
                timeout=120,
                errors="replace",
            )
            if r.returncode == 0 and (r.stdout or "").strip():
                return r.stdout.strip()
        except Exception:
            pass
    return None


def extract_text_from_any_file(file_path):
    """Detects file type and extracts text accordingly."""
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    frappe.logger().info(f"Starting extraction for file: {file_path} with extension: {ext}")

    if ext == ".pdf":
        text = extract_text(file_path)
        if not text.strip():
            # Image-based PDF: return None - caller will use parse_with_gemini_file (direct PDF to Gemini)
            # This bypasses OCR/poppler which often fails and is slow
            return None
    
    elif ext == ".docx":
        doc = docx.Document(file_path)
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(parts)

    elif ext == ".doc":
        extracted = extract_text_from_legacy_doc(file_path)
        if not extracted:
            frappe.logger().warning(f"Could not extract text from legacy .doc: {file_path}")
            return None
        text = extracted
    
    elif ext in [".jpg", ".jpeg", ".png"]:
        text = pytesseract.image_to_string(Image.open(file_path))
    
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    
    else:
        frappe.logger().warning(f"Unsupported file format: {ext}")
        return None

    return text.strip()

def extract_text_with_ocr(file_path):
    """OCR for Scanned PDFs."""
    try:
        images = convert_from_path(file_path)
        return "\n".join([pytesseract.image_to_string(img) for img in images])
    except Exception as e:
        frappe.logger().error(f"OCR Error: {str(e)}")
        return ""

import base64
def parse_with_gemini_file(file_path, job_title=None, job_description=None):
    """Send PDF directly to Gemini - bypasses OCR, works for image-based PDFs. Faster."""
    api_key = frappe.conf.get("gemini_api_key")
    if not api_key:
        raise Exception("Gemini API key missing in site_config.json")

    genai.configure(api_key=api_key)
    prompt_path = frappe.get_app_path("resume", "resume", "doctype", "pdf_upload", "resume_prompt.txt")
    with open(prompt_path, "r") as f:
        prompt_template = f.read()

    prompt = prompt_template.replace("{{RESUME_TEXT}}", "Resume attached as PDF.")
    prompt = prompt.replace("{{JOB_TITLE}}", job_title or "N/A")
    prompt = prompt.replace("{{JOB_DESCRIPTION}}", job_description or "N/A")

    with open(file_path, "rb") as f:
        pdf_bytes = f.read()

    model_names = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-2.5-pro"]  # flash first for speed
    last_error = None
    for model_name in model_names:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(
                [
                    {"mime_type": "application/pdf", "data": pdf_bytes},
                    prompt
                ]
            )
            text = response.text.strip()
            if text.startswith("```"):
                text = text.replace("```json", "").replace("```", "").strip()
            return json.loads(text)
        except Exception as e:
            last_error = e
            frappe.logger().warning(f"Gemini {model_name} failed for PDF: {e}")
            continue
    raise last_error

# --- Background Job (Parallel Processing) ---

def _is_retryable_error(exc):
    """Check if error is rate limit or transient (worth retrying)."""
    msg = str(exc).lower()
    return any(x in msg for x in ["429", "resource_exhausted", "rate limit", "quota", "503", "overloaded"])

def _call_gemini_with_retry(call_fn):
    """Execute Gemini API call with exponential backoff on rate limits."""
    last_error = None
    for attempt in range(GEMINI_RETRY_ATTEMPTS):
        try:
            return call_fn()
        except Exception as e:
            last_error = e
            if attempt < GEMINI_RETRY_ATTEMPTS - 1 and _is_retryable_error(e):
                delay = GEMINI_RETRY_BASE_DELAY * (2 ** attempt)
                time.sleep(delay)
            else:
                raise
    raise last_error

def _parse_text_threadsafe(api_key, prompt_template, text, job_title, job_desc):
    """Thread-safe: parse text with Gemini. Retries on rate limit."""
    genai.configure(api_key=api_key)
    prompt = prompt_template.replace("{{RESUME_TEXT}}", text)
    prompt = prompt.replace("{{JOB_TITLE}}", job_title or "N/A")
    prompt = prompt.replace("{{JOB_DESCRIPTION}}", job_desc or "N/A")
    last_error = None
    for model_name in ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-2.5-pro"]:
        try:
            def _call(m=model_name, p=prompt):
                model = genai.GenerativeModel(m)
                response = model.generate_content(p)
                cleaned = response.text.strip()
                if cleaned.startswith("```json"): cleaned = cleaned[7:]
                if cleaned.endswith("```"): cleaned = cleaned[:-3]
                return json.loads(cleaned.strip())
            return _call_gemini_with_retry(lambda: _call())
        except Exception as e:
            last_error = e
            continue
    raise last_error or Exception("All Gemini models failed")

def _parse_pdf_threadsafe(api_key, prompt_template, file_path, job_title, job_desc):
    """Thread-safe: send PDF to Gemini. Retries on rate limit."""
    genai.configure(api_key=api_key)
    prompt = prompt_template.replace("{{RESUME_TEXT}}", "Resume attached as PDF.")
    prompt = prompt.replace("{{JOB_TITLE}}", job_title or "N/A")
    prompt = prompt.replace("{{JOB_DESCRIPTION}}", job_desc or "N/A")
    with open(file_path, "rb") as f:
        pdf_bytes = f.read()
    last_error = None
    for model_name in ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-2.5-pro"]:
        try:
            def _call(m=model_name, p=prompt, b=pdf_bytes):
                model = genai.GenerativeModel(m)
                response = model.generate_content([{"mime_type": "application/pdf", "data": b}, p])
                text = response.text.strip()
                if text.startswith("```"):
                    text = text.replace("```json", "").replace("```", "").strip()
                return json.loads(text)
            return _call_gemini_with_retry(lambda: _call())
        except Exception as e:
            last_error = e
            continue
    raise last_error or Exception("All Gemini models failed for PDF")

def _extract_text(file_path, ext):
    """Extract text from file. No frappe logging."""
    if ext == ".pdf":
        text = extract_text(file_path)
        return text.strip() if text else None
    elif ext == ".docx":
        doc = docx.Document(file_path)
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)
    elif ext == ".doc":
        return extract_text_from_legacy_doc(file_path)
    elif ext in [".jpg", ".jpeg", ".png"]:
        return pytesseract.image_to_string(Image.open(file_path))
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    return None

def _extract_and_parse_file(args):
    """Runs in thread. Uses pre-loaded api_key + prompt_template. No frappe calls."""
    file_path, file_url, job_title, job_desc, ext, api_key, prompt_template = args
    try:
        if ext == ".pdf":
            text = _extract_text(file_path, ext)
            if text:
                applicant_data = _parse_text_threadsafe(api_key, prompt_template, text, job_title, job_desc)
            else:
                applicant_data = _parse_pdf_threadsafe(api_key, prompt_template, file_path, job_title, job_desc)
        else:
            text = _extract_text(file_path, ext)
            if not text:
                return (file_url, None, f"Could not extract text from {file_url}")
            applicant_data = _parse_text_threadsafe(api_key, prompt_template, text, job_title, job_desc)
        return (file_url, applicant_data, None)
    except Exception as e:
        return (file_url, None, str(e))


def process_files_background(docname):
    """Processes uploaded files in parallel - multiple resumes parsed concurrently for speed."""
    t_start = time.perf_counter()
    try:
        doc = frappe.get_doc("PDF Upload", docname)
        if not doc.get("files"):
            msg = f"PDF Upload {docname}: No rows in Files table."
            frappe.logger().warning(msg)
            frappe.log_error(msg, "PDF Upload — no files")
            return

        job_desc = frappe.db.get_value("Job Opening", doc.job_title, "description") if doc.job_title else None
        job_title = doc.job_title

        if doc.job_title:
            jo_status = frappe.db.get_value("Job Opening", doc.job_title, "status")
            if jo_status == "Closed":
                msg = _(
                    "Skipped processing: Job Opening {0} is Closed. Applicants cannot be created until the opening is reopened or you change the Job Opening on this PDF Upload."
                ).format(doc.job_title)
                frappe.log_error(title="PDF Upload — Job Opening closed", message=msg)
                try:
                    doc.add_comment("Comment", text=msg)
                    frappe.db.commit()
                except Exception:
                    pass
                return

        # Pre-load for threads (no frappe in worker threads)
        api_key = frappe.conf.get("gemini_api_key")
        if not api_key:
            msg = "PDF Upload resume job skipped: gemini_api_key missing in site_config.json"
            frappe.logger().error(msg)
            frappe.log_error(msg, "PDF Upload — Gemini not configured")
            return
        prompt_path = frappe.get_app_path("resume", "resume", "doctype", "pdf_upload", "resume_prompt.txt")
        with open(prompt_path, "r") as f:
            prompt_template = f.read()

        # 1. Resolve file paths (fast, sequential)
        t_path_start = time.perf_counter()
        tasks = []
        for file_entry in doc.get("files"):
            file_url = file_entry.get("file_upload")
            if not file_url:
                continue
            file_url = normalize_attach_file_url(file_url)

            file_path = None
            file_list = frappe.get_all("File", filters={"file_url": file_url}, limit=1)
            if file_list:
                try:
                    file_doc = frappe.get_doc("File", file_list[0].name)
                    file_path = file_doc.get_full_path()
                except Exception:
                    pass
            if not file_path or not os.path.exists(file_path):
                path_parts = file_url.lstrip("/").split("/") if file_url.startswith("/private") else ["public"] + file_url.lstrip("/").split("/")
                file_path = os.path.abspath(frappe.get_site_path(*path_parts))

            if not os.path.exists(file_path):
                frappe.logger().error(f"File not found on server: {file_path}")
                continue

            ext = os.path.splitext(file_path)[1].lower()
            tasks.append((file_path, file_url, job_title, job_desc, ext, api_key, prompt_template))

        if not tasks:
            frappe.log_error(
                f"PDF Upload {docname}: No file could be read from disk (missing File row, wrong URL, or path).",
                "PDF Upload — file path resolution failed",
            )
            return

        t_path_elapsed = time.perf_counter() - t_path_start
        frappe.logger().info(f"PDF Upload {docname}: Path resolution {t_path_elapsed:.2f}s for {len(tasks)} file(s)")

        # 2. Parse files (2 at a time to avoid rate limits; retries on 429)
        t_parse_start = time.perf_counter()
        workers = min(PARALLEL_WORKERS, len(tasks))
        results = []
        parse_errors = []
        with ThreadPoolExecutor(max_workers=workers) as executor:
            futures = {}
            for i, t in enumerate(tasks):
                if i > 0:
                    time.sleep(2)  # Stagger starts to reduce rate limit spikes
                futures[executor.submit(_extract_and_parse_file, t)] = t
            for future in as_completed(futures):
                file_url, applicant_data, error_msg = future.result()
                if error_msg:
                    frappe.logger().error(f"Parsing failed for {file_url}: {error_msg}")
                    parse_errors.append(f"{file_url}: {error_msg}")
                elif applicant_data:
                    results.append((file_url, applicant_data))

        t_parse_elapsed = time.perf_counter() - t_parse_start
        frappe.logger().info(f"PDF Upload {docname}: AI parsing {t_parse_elapsed:.2f}s ({len(results)}/{len(tasks)} parsed)")

        # 3. Create Job Applicants (fast, sequential - DB writes)
        t_db_start = time.perf_counter()
        status_options = frappe.get_meta("Job Applicant").get_options("status") or ""
        valid_statuses = [s.strip() for s in status_options.split("\n") if s.strip()]
        default_status = "CV Submitted" if "CV Submitted" in valid_statuses else (valid_statuses[0] if valid_statuses else "Open")

        created_names = []
        skipped_duplicate = 0
        skipped_missing = 0
        insert_failed = []

        for file_url, applicant_data in results:
            email = applicant_data.get("email_id")
            if not email or not applicant_data.get("applicant_name"):
                frappe.logger().warning(f"Missing Name/Email in AI response for {file_url}")
                skipped_missing += 1
                continue

            if frappe.db.exists("Job Applicant", {"email_id": email}):
                frappe.logger().info(f"Skipping: Applicant {email} already exists.")
                skipped_duplicate += 1
                continue

            try:
                new_applicant = frappe.get_doc({
                    "doctype": "Job Applicant",
                    "applicant_name": applicant_data.get("applicant_name"),
                    "email_id": email,
                    "phone_number": applicant_data.get("phone_number"),
                    "custom_phone_number_2": applicant_data.get("custom_phone_number_2", ""),
                    "resume_attachment": file_url,
                    "status": default_status,
                    "job_title": doc.job_title,
                    "designation": doc.designation,
                    "applicant_rating": applicant_data.get("applicant_rating"),
                    "score": applicant_data.get("score"),
                    "fit_level": applicant_data.get("fit_level"),
                    "justification_by_ai": applicant_data.get("justification_by_ai"),
                    "custom_yes": 1
                })
                new_applicant.insert(ignore_permissions=True)
                created_names.append(new_applicant.name)
                frappe.logger().info(f"Successfully Created Applicant: {new_applicant.name} from {file_url}")
            except Exception as e:
                frappe.logger().error(f"Database insertion failed for {email}: {str(e)}")
                frappe.log_error(
                    frappe.get_traceback(with_context=True),
                    f"PDF Upload Job Applicant insert failed ({email})",
                )
                insert_failed.append(f"{email}: {e!s}")

        t_db_elapsed = time.perf_counter() - t_db_start
        t_total = time.perf_counter() - t_start
        frappe.logger().info(
            f"PDF Upload {docname}: DB writes {t_db_elapsed:.2f}s | Total {t_total:.2f}s "
            f"(path:{t_path_elapsed:.1f}s parse:{t_parse_elapsed:.1f}s db:{t_db_elapsed:.1f}s)"
        )

        # Visible summary on the PDF Upload form (timeline)
        summary_parts = [
            _("Resume import finished in {0:.1f}s.").format(t_total),
            _("Created: {0}").format(len(created_names)),
            _("Skipped (duplicate email): {0}").format(skipped_duplicate),
            _("Skipped (missing name/email from AI): {0}").format(skipped_missing),
            _("Parse errors: {0}").format(len(parse_errors)),
            _("Insert errors: {0}").format(len(insert_failed)),
        ]
        if created_names:
            summary_parts.append(_("Applicants: {0}").format(", ".join(created_names[:20])))
        if parse_errors:
            summary_parts.append(_("First error: {0}").format(parse_errors[0][:500]))
        try:
            doc.add_comment("Comment", text="\n".join(summary_parts))
        except Exception:
            frappe.logger().error(f"PDF Upload {docname}: could not add summary comment")

        frappe.db.commit()

    except Exception as e:
        frappe.logger().error(f"Critical error in background job {docname}: {str(e)}")
        frappe.log_error(frappe.get_traceback(with_context=True), f"PDF Upload critical {docname}")
        


#####
def process_resume_pipeline(file_path, job_title=None, job_desc=None):
    """
    Step 1: Try normal extraction
    Step 2: If failed → OCR
    Step 3: If failed → Gemini fallback
    """

    ext = os.path.splitext(file_path)[1].lower()

    # ---------------- STEP 1: NORMAL EXTRACTION ----------------
    text = extract_text_from_any_file(file_path)
    
    frappe.log_error(
        title="process_resume_pipeline - extracted text",
        message=f"Extracted text {text}"   ) 

    # ---------------- STEP 2: OCR ----------------
    if not text or len(text.strip()) < 30:
        frappe.logger().info("⚠️ Low text detected → Trying OCR")

        try:
            text = extract_text_with_ocr(file_path)
        except Exception as e:
            frappe.logger().error(f"OCR failed: {e}")
            text = None

    # ---------------- STEP 3: GEMINI FALLBACK ----------------
    if not text or len(text.strip()) < 30:
        frappe.logger().info("🚀 Falling back to Gemini")

        try:
            # If PDF → send file directly (BEST)
            if ext == ".pdf":
                return parse_with_gemini_file(file_path, job_title, job_desc)

            # Otherwise send extracted text (even if partial)
            return parse_with_gemini(text or "", job_title, job_desc)

        except Exception as e:
            frappe.logger().error(f"Gemini failed: {e}")
            return None

    # ---------------- STEP 4: PARSE TEXT VIA GEMINI ----------------
    try:
        return parse_with_gemini(text, job_title, job_desc)
    except Exception as e:
        frappe.logger().error(f"Gemini parse failed: {e}")
        return None